import argparse
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path

import cv2
import numpy as np
import pytesseract
import torch
from docx import Document
from spellchecker import SpellChecker

from character_cnn import CharacterCNN, CnnConfig

DEFAULT_WINDOWS_TESSERACT_PATHS = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
)
DEFAULT_CNN_CHECKPOINT = Path("models/character_cnn.pt")


@dataclass
class OcrCandidate:
    text: str
    score: float
    average_confidence: float
    strategy: str
    words: list["OcrWord"] = field(default_factory=list)


@dataclass(frozen=True)
class OcrWord:
    text: str
    confidence: float
    bbox: tuple[int, int, int, int]
    block_num: int
    par_num: int
    line_num: int
    word_num: int


@dataclass(frozen=True)
class CnnBundle:
    model: CharacterCNN
    class_names: list[str]
    image_size: int
    device: torch.device


@dataclass(frozen=True)
class CnnRefinementStats:
    total_alpha_words: int
    words_with_cnn: int
    words_replaced: int


def configure_tesseract() -> None:
    if shutil.which("tesseract"):
        return
    for path in DEFAULT_WINDOWS_TESSERACT_PATHS:
        if Path(path).exists():
            pytesseract.pytesseract.tesseract_cmd = path
            return


def deskew(gray_image: np.ndarray) -> np.ndarray:
    _, thresh = cv2.threshold(
        gray_image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
    )
    coords = np.column_stack(np.where(thresh > 0))
    if coords.size == 0:
        return gray_image

    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle

    if abs(angle) < 0.25:
        return gray_image

    height, width = gray_image.shape[:2]
    matrix = cv2.getRotationMatrix2D((width // 2, height // 2), angle, 1.0)
    return cv2.warpAffine(
        gray_image,
        matrix,
        (width, height),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_REPLICATE,
    )


def preprocess_image_for_ocr(image_path: Path) -> dict[str, np.ndarray]:
    image = cv2.imread(str(image_path))
    if image is None:
        raise FileNotFoundError(f"Could not read image: {image_path}")

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    upscaled = cv2.resize(gray, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
    contrast = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(upscaled)
    denoised = cv2.GaussianBlur(contrast, (3, 3), 0)
    aligned = deskew(denoised)

    _, otsu = cv2.threshold(aligned, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    adaptive = cv2.adaptiveThreshold(
        aligned,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        11,
    )
    adaptive_inv = cv2.bitwise_not(adaptive)
    return {
        "gray": aligned,
        "otsu": otsu,
        "adaptive": adaptive,
        "adaptive_inv": adaptive_inv,
    }


def parse_ocr_words(data: dict[str, list[str]]) -> list[OcrWord]:
    words: list[OcrWord] = []
    texts = data.get("text", [])
    confs = data.get("conf", [])
    lefts = data.get("left", [])
    tops = data.get("top", [])
    widths = data.get("width", [])
    heights = data.get("height", [])
    blocks = data.get("block_num", [])
    pars = data.get("par_num", [])
    lines = data.get("line_num", [])
    word_nums = data.get("word_num", [])

    for idx, raw_text in enumerate(texts):
        text = (raw_text or "").strip()
        if not text:
            continue

        try:
            conf = float(confs[idx])
        except (TypeError, ValueError, IndexError):
            conf = -1.0
        if conf < 0:
            continue

        try:
            word = OcrWord(
                text=text,
                confidence=conf,
                bbox=(
                    int(lefts[idx]),
                    int(tops[idx]),
                    int(widths[idx]),
                    int(heights[idx]),
                ),
                block_num=int(blocks[idx]),
                par_num=int(pars[idx]),
                line_num=int(lines[idx]),
                word_num=int(word_nums[idx]),
            )
        except (ValueError, IndexError):
            continue
        words.append(word)
    return words


def score_candidate(text: str, confidences: list[float]) -> float:
    avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
    alpha_chars = sum(ch.isalpha() for ch in text)
    alpha_ratio = alpha_chars / max(1, len(text))
    return avg_conf + (alpha_ratio * 14.0) + min(len(text.split()), 30) * 0.25


def extract_text(preprocessed_variants: dict[str, np.ndarray]) -> tuple[str, OcrCandidate]:
    configure_tesseract()
    candidates: list[OcrCandidate] = []

    for name, image in preprocessed_variants.items():
        for psm in (6, 4, 11):
            config = f"--oem 3 --psm {psm}"
            raw_text = pytesseract.image_to_string(image, config=config).strip()
            if not raw_text:
                continue

            data = pytesseract.image_to_data(
                image,
                config=config,
                output_type=pytesseract.Output.DICT,
            )
            confidences = []
            for conf in data.get("conf", []):
                try:
                    value = float(conf)
                except (TypeError, ValueError):
                    continue
                if value >= 0:
                    confidences.append(value)

            score = score_candidate(raw_text, confidences)
            candidate = OcrCandidate(
                text=raw_text,
                score=score,
                average_confidence=(sum(confidences) / len(confidences)) if confidences else 0.0,
                strategy=f"{name}:psm{psm}",
                words=parse_ocr_words(data),
            )
            candidates.append(candidate)

    if not candidates:
        raise RuntimeError(
            "OCR produced no text. Verify that Tesseract is installed and the image is readable."
        )

    best = max(candidates, key=lambda candidate: candidate.score)
    return best.text, best


def edit_distance(lhs: str, rhs: str) -> int:
    if lhs == rhs:
        return 0
    if not lhs:
        return len(rhs)
    if not rhs:
        return len(lhs)

    previous = list(range(len(rhs) + 1))
    for i, ch_l in enumerate(lhs, start=1):
        current = [i]
        for j, ch_r in enumerate(rhs, start=1):
            cost = 0 if ch_l == ch_r else 1
            current.append(
                min(
                    previous[j] + 1,
                    current[j - 1] + 1,
                    previous[j - 1] + cost,
                )
            )
        previous = current
    return previous[-1]


def normalize_ocr_token(token: str) -> str:
    if any(ch.isalpha() for ch in token) and any(ch.isdigit() for ch in token):
        return token.translate(str.maketrans({"0": "o", "1": "l", "5": "s"}))
    return token


def decode_class_name(class_name: str) -> str:
    if len(class_name) == 1:
        return class_name

    normalized = class_name.lower()
    symbol_map = {
        "space": " ",
        "dot": ".",
        "period": ".",
        "comma": ",",
        "apostrophe": "'",
        "quote": '"',
        "question": "?",
        "exclamation": "!",
        "dash": "-",
        "minus": "-",
    }
    if normalized in symbol_map:
        return symbol_map[normalized]

    for prefix in ("upper_", "lower_", "digit_", "char_"):
        if normalized.startswith(prefix):
            suffix = normalized.split(prefix, 1)[1]
            if len(suffix) == 1:
                return suffix

    if len(normalized) == 1:
        return normalized
    return ""


def prepare_character_tensor(character_binary: np.ndarray, image_size: int) -> torch.Tensor:
    ys, xs = np.where(character_binary > 0)
    if len(xs) == 0 or len(ys) == 0:
        canvas = np.zeros((image_size, image_size), dtype=np.uint8)
    else:
        x0, x1 = xs.min(), xs.max() + 1
        y0, y1 = ys.min(), ys.max() + 1
        crop = character_binary[y0:y1, x0:x1]

        height, width = crop.shape
        side = max(height, width) + 8
        square = np.zeros((side, side), dtype=np.uint8)
        y_offset = (side - height) // 2
        x_offset = (side - width) // 2
        square[y_offset : y_offset + height, x_offset : x_offset + width] = crop
        canvas = cv2.resize(square, (image_size, image_size), interpolation=cv2.INTER_AREA)

    tensor = torch.from_numpy(canvas.astype(np.float32) / 255.0)
    return tensor.unsqueeze(0).unsqueeze(0)


def infer_word_with_cnn(word_image: np.ndarray, bundle: CnnBundle) -> tuple[str, float]:
    if word_image.size == 0:
        return "", 0.0

    if len(word_image.shape) == 3:
        gray = cv2.cvtColor(word_image, cv2.COLOR_BGR2GRAY)
    else:
        gray = word_image

    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    if float(np.mean(binary)) > 127:
        binary = cv2.bitwise_not(binary)

    cleaned = cv2.morphologyEx(binary, cv2.MORPH_OPEN, np.ones((2, 2), np.uint8))
    contours, _ = cv2.findContours(cleaned, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return "", 0.0

    height, width = cleaned.shape
    min_area = max(12, int(0.001 * height * width))
    boxes: list[tuple[int, int, int, int]] = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        area = w * h
        if area < min_area or h < 5:
            continue
        boxes.append((x, y, w, h))
    boxes.sort(key=lambda box: box[0])

    if not boxes:
        return "", 0.0

    chars: list[str] = []
    confidences: list[float] = []
    with torch.no_grad():
        for x, y, w, h in boxes:
            char_crop = cleaned[y : y + h, x : x + w]
            tensor = prepare_character_tensor(char_crop, bundle.image_size).to(bundle.device)
            logits = bundle.model(tensor)
            probabilities = torch.softmax(logits, dim=1)
            confidence, index = torch.max(probabilities, dim=1)

            class_name = bundle.class_names[index.item()]
            decoded = decode_class_name(class_name)
            if not decoded:
                continue
            chars.append(decoded)
            confidences.append(float(confidence.item()))

    if not chars:
        return "", 0.0
    return "".join(chars), (sum(confidences) / len(confidences))


def apply_case(template: str, word: str) -> str:
    if template.isupper():
        return word.upper()
    if len(template) > 1 and template[0].isupper() and template[1:].islower():
        return word.capitalize()
    return word


def choose_word_from_ocr_and_cnn(
    ocr_word: str,
    ocr_confidence: float,
    cnn_word: str,
    cnn_confidence: float,
    spell: SpellChecker,
) -> str:
    if not cnn_word:
        return ocr_word

    ocr_norm = ocr_word.lower()
    cnn_norm = cnn_word.lower()
    if ocr_norm == cnn_norm:
        return ocr_word

    ocr_in_dict = ocr_norm in spell
    cnn_in_dict = cnn_norm in spell

    if cnn_in_dict and not ocr_in_dict and cnn_confidence >= 0.40:
        return apply_case(ocr_word, cnn_word)
    if ocr_in_dict and not cnn_in_dict and ocr_confidence >= 45.0:
        return ocr_word

    distance = edit_distance(ocr_norm, cnn_norm)
    if distance <= 1 and cnn_confidence >= 0.30:
        return apply_case(ocr_word, cnn_word)
    if cnn_in_dict and cnn_confidence >= 0.55:
        return apply_case(ocr_word, cnn_word)
    if cnn_confidence >= 0.78 and distance <= max(2, len(ocr_norm) // 2):
        return apply_case(ocr_word, cnn_word)

    return ocr_word


def build_text_from_words(words: list[tuple[OcrWord, str]]) -> str:
    if not words:
        return ""

    ordered = sorted(
        words,
        key=lambda item: (
            item[0].block_num,
            item[0].par_num,
            item[0].line_num,
            item[0].word_num,
        ),
    )

    lines: list[str] = []
    current_key: tuple[int, int, int] | None = None
    current_words: list[str] = []
    for ocr_word, text in ordered:
        line_key = (ocr_word.block_num, ocr_word.par_num, ocr_word.line_num)
        if current_key is None:
            current_key = line_key
        if line_key != current_key:
            lines.append(" ".join(current_words))
            current_words = []
            current_key = line_key
        current_words.append(text)

    if current_words:
        lines.append(" ".join(current_words))
    return "\n".join(lines)


def refine_text_with_cnn(
    image_for_cnn: np.ndarray,
    ocr_words: list[OcrWord],
    bundle: CnnBundle,
) -> tuple[str, CnnRefinementStats]:
    spell = SpellChecker(distance=2)
    refined: list[tuple[OcrWord, str]] = []

    total_alpha_words = 0
    words_with_cnn = 0
    words_replaced = 0

    image_h, image_w = image_for_cnn.shape[:2]
    for word in ocr_words:
        replacement = word.text
        if re.fullmatch(r"[A-Za-z]+", word.text):
            total_alpha_words += 1
            x, y, w, h = word.bbox
            x0 = max(0, x - 2)
            y0 = max(0, y - 2)
            x1 = min(image_w, x + w + 2)
            y1 = min(image_h, y + h + 2)
            crop = image_for_cnn[y0:y1, x0:x1]

            cnn_word, cnn_confidence = infer_word_with_cnn(crop, bundle)
            if cnn_word:
                words_with_cnn += 1
                replacement = choose_word_from_ocr_and_cnn(
                    word.text,
                    word.confidence,
                    cnn_word,
                    cnn_confidence,
                    spell,
                )
                if replacement.lower() != word.text.lower():
                    words_replaced += 1

        refined.append((word, replacement))

    return build_text_from_words(refined), CnnRefinementStats(
        total_alpha_words=total_alpha_words,
        words_with_cnn=words_with_cnn,
        words_replaced=words_replaced,
    )


def load_cnn_bundle(checkpoint_path: Path | None) -> CnnBundle | None:
    if checkpoint_path is None:
        return None
    if not checkpoint_path.exists():
        return None

    payload = torch.load(str(checkpoint_path), map_location="cpu")
    class_names = payload.get("class_names")
    state_dict = payload.get("model_state_dict")
    if not class_names or state_dict is None:
        return None

    image_size = int(payload.get("image_size", 64))
    num_classes = int(payload.get("num_classes", len(class_names)))
    if num_classes != len(class_names):
        num_classes = len(class_names)

    model = CharacterCNN(num_classes=num_classes, config=CnnConfig(image_size=image_size))
    model.load_state_dict(state_dict)
    model.eval()

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)

    return CnnBundle(
        model=model,
        class_names=[str(name) for name in class_names],
        image_size=image_size,
        device=device,
    )


def choose_best_correction(word: str, candidates: set[str], spell: SpellChecker) -> str:
    if not candidates:
        return word

    filtered = list(candidates)
    if len(word) >= 3:
        same_edges = [c for c in filtered if c[0] == word[0] and c[-1] == word[-1]]
        if same_edges:
            filtered = same_edges

    return min(
        filtered,
        key=lambda candidate: (
            edit_distance(word, candidate),
            -spell.word_frequency.dictionary.get(candidate, 0),
            abs(len(candidate) - len(word)),
        ),
    )


def predict_words(raw_text: str) -> str:
    spell = SpellChecker(distance=2)
    token_pattern = re.compile(r"[A-Za-z']+|[^A-Za-z']+")
    tokens = token_pattern.findall(raw_text)

    corrected = []
    for token in tokens:
        if not re.fullmatch(r"[A-Za-z']+", token):
            corrected.append(token)
            continue

        # Preserve very short words and likely acronyms.
        if len(token) <= 2 or token.isupper():
            corrected.append(token)
            continue

        normalized = normalize_ocr_token(token.lower())
        if normalized in spell:
            corrected.append(token)
            continue

        candidate_words = spell.candidates(normalized) or set()
        replacement = choose_best_correction(normalized, candidate_words, spell)

        # Guardrail: avoid aggressive substitutions that change too many characters.
        if edit_distance(normalized, replacement) > max(2, len(normalized) // 2):
            corrected.append(token)
            continue

        if token[0].isupper():
            replacement = replacement.capitalize()
        corrected.append(replacement)

    return "".join(corrected)


def save_outputs(predicted_text: str, output_stem: Path) -> tuple[Path, Path]:
    txt_path = output_stem.with_suffix(".txt")
    docx_path = output_stem.with_suffix(".docx")

    txt_path.parent.mkdir(parents=True, exist_ok=True)
    txt_path.write_text(predicted_text, encoding="utf-8")

    document = Document()
    document.add_heading("Predicted Handwriting Text", level=1)
    document.add_paragraph(predicted_text)
    document.save(str(docx_path))

    return txt_path, docx_path


def run_pipeline(
    image_path: Path,
    output_stem: Path,
    cnn_checkpoint: Path | None = DEFAULT_CNN_CHECKPOINT,
) -> tuple[str, str, Path, Path]:
    preprocessed = preprocess_image_for_ocr(image_path)
    raw_text, best_candidate = extract_text(preprocessed)

    combined_text = raw_text
    cnn_bundle = load_cnn_bundle(cnn_checkpoint)
    if cnn_bundle and best_candidate.words:
        cnn_assisted_text, _ = refine_text_with_cnn(
            preprocessed["adaptive_inv"],
            best_candidate.words,
            cnn_bundle,
        )
        if cnn_assisted_text.strip():
            combined_text = cnn_assisted_text

    predicted_text = predict_words(combined_text)
    txt_path, docx_path = save_outputs(predicted_text, output_stem)
    return raw_text, predicted_text, txt_path, docx_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Read a handwritten image, recognize letters/words, apply "
            "word prediction/correction, and export output documents."
        )
    )
    parser.add_argument(
        "--image",
        required=True,
        help="Path to the input handwriting image file.",
    )
    parser.add_argument(
        "--output-stem",
        default="outputs/predicted_document",
        help="Output file stem (without extension).",
    )
    parser.add_argument(
        "--cnn-checkpoint",
        default=str(DEFAULT_CNN_CHECKPOINT),
        help="Path to CharacterCNN checkpoint saved by Letter_Detection.py.",
    )
    parser.add_argument(
        "--disable-cnn",
        action="store_true",
        help="Disable CharacterCNN-assisted refinement and use OCR only.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    image_path = Path(args.image)
    output_stem = Path(args.output_stem)

    cnn_checkpoint = None if args.disable_cnn else Path(args.cnn_checkpoint)
    raw_text, predicted_text, txt_path, docx_path = run_pipeline(
        image_path,
        output_stem,
        cnn_checkpoint=cnn_checkpoint,
    )

    print("Raw OCR text:")
    print(raw_text.strip())
    print("\nPredicted text:")
    print(predicted_text.strip())
    print(f"\nSaved text output: {txt_path}")
    print(f"Saved document output: {docx_path}")


if __name__ == "__main__":
    main()
